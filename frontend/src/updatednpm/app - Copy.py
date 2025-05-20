from fastapi import FastAPI, Query, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from transformers import (
    MarianMTModel, MarianTokenizer, 
    AutoTokenizer, AutoModelForSeq2SeqLM,
    LayoutLMv3Processor, LayoutLMv3ForTokenClassification
)
from PIL import Image
import torch
import nltk
from docx import Document
from typing import Annotated, Optional
import os
from enum import Enum
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Download the NLTK sentence tokenizer data
nltk.download('punkt')

# Language support
class Language(str, Enum):
    ENGLISH = "en"
    TAMIL = "ta"
    TELUGU = "te"
    MALAYALAM = "ml"
    KANNADA = "kn"
    FRENCH = "fr"
    SPANISH = "es"

# Model configurations
MODEL_CONFIG = {
    "en-ta": "suriya7/English-to-Tamil",
    "en-te": "Helsinki-NLP/opus-mt-en-mul",  # Will filter for Telugu
    "en-ml": "Helsinki-NLP/opus-mt-en-mul",  # Will filter for Malayalam
    "en-kn": "Helsinki-NLP/opus-mt-en-mul",  # Will filter for Kannada
    "default": "Helsinki-NLP/opus-mt-en-mul"
}

# Load layout-aware model for document structure preservation
def load_layout_model():
    processor = LayoutLMv3Processor.from_pretrained("microsoft/layoutlmv3-base")
    model = LayoutLMv3ForTokenClassification.from_pretrained("microsoft/layoutlmv3-base")
    return processor, model

# Optimized model loading with quantization
def load_model(from_lang: str, to_lang: str, quantize: bool = True):
    model_key = f"{from_lang}-{to_lang}"
    model_name = MODEL_CONFIG.get(model_key, MODEL_CONFIG["default"])
    
    logger.info(f"Loading model: {model_name}")
    
    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
    
    if quantize:
        # Apply dynamic quantization for faster inference
        model = torch.quantization.quantize_dynamic(
            model, {torch.nn.Linear}, dtype=torch.qint8
        )
        logger.info("Model quantization applied")
    
    return model, tokenizer

def process_document_layout(file_path: str):
    """Process document to preserve layout (tables, images)"""
    try:
        if file_path.endswith('.pdf'):
            # Use OCR and layout analysis for PDFs
            images = convert_pdf_to_images(file_path)
            processor, model = load_layout_model()
            
            layout_info = []
            for image in images:
                encoding = processor(image, return_tensors="pt")
                outputs = model(**encoding)
                layout_info.append(outputs)
            
            return layout_info
        elif file_path.endswith(('.docx', '.doc')):
            # Process Word document structure
            doc = Document(file_path)
            layout_info = []
            for para in doc.paragraphs:
                layout_info.append({
                    'text': para.text,
                    'style': para.style.name,
                    'runs': [run.text for run in para.runs]
                })
            return layout_info
    except Exception as e:
        logger.error(f"Layout processing error: {str(e)}")
        return None

@app.post("/translate_document")
async def translate_document(
    file: UploadFile = File(...),
    from_lang: Language = Form(Language.ENGLISH),
    to_lang: Language = Form(Language.TAMIL),
    domain: Optional[str] = Form(None),
    preserve_layout: bool = Form(False)
):
    """Enhanced document translation endpoint with layout preservation"""
    try:
        # Save uploaded file temporarily
        temp_path = f"temp_{file.filename}"
        with open(temp_path, "wb") as temp_file:
            temp_file.write(await file.read())
        
        # Process document layout if requested
        layout_info = None
        if preserve_layout:
            layout_info = process_document_layout(temp_path)
        
        # Read document content
        if temp_path.endswith('.docx'):
            doc = Document(temp_path)
            text = "\n".join([para.text for para in doc.paragraphs])
        elif temp_path.endswith('.txt'):
            with open(temp_path, 'r') as f:
                text = f.read()
        else:
            return {"error": "Unsupported file type"}
        
        # Load appropriate model
        model, tokenizer = load_model(from_lang.value, to_lang.value)
        
        # Batch processing for large documents
        sentences = nltk.sent_tokenize(text)
        batch_size = 8  # Optimal batch size for memory efficiency
        translated_batches = []
        
        for i in range(0, len(sentences), batch_size):
            batch = sentences[i:i + batch_size]
            inputs = tokenizer(
                batch, 
                return_tensors='pt', 
                padding=True, 
                truncation=True
            )
            with torch.no_grad():
                outputs = model.generate(**inputs)
            translated = tokenizer.batch_decode(outputs, skip_special_tokens=True)
            translated_batches.extend(translated)
        
        translated_text = " ".join(translated_batches)
        
        # Clean up
        os.remove(temp_path)
        
        return {
            "from": from_lang.value,
            "to": to_lang.value,
            "translated_text": translated_text,
            "layout_info": layout_info if preserve_layout else None
        }
        
    except Exception as e:
        logger.error(f"Translation error: {str(e)}")
        return {"error": str(e)}

@app.get("/supported_languages")
async def get_supported_languages():
    """Endpoint to list supported languages"""
    return {
        "languages": [
            {"code": lang.value, "name": lang.name}
            for lang in Language
        ]
    }

# Additional endpoints for domain adaptation, etc. would go here